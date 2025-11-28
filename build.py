# This is the master script to build provincially-configured executables.
# Run this script from your terminal: python build.py

import os
import sys
import shutil
import subprocess
import tempfile
import time
import traceback
import datetime
from pathlib import Path
import glob
import ctypes
from datetime import datetime
import re
import warnings

# Suppress specific deprecation warnings
warnings.filterwarnings(
    'ignore',
    message='pkg_resources is deprecated as an API',
    category=UserWarning,
    module='docxcompose.properties'
)

try:
    from docx import Document
    from docxtpl import DocxTemplate
    from docx2pdf import convert
except ImportError as e:
    print("Error: Required packages not found. Please install them using:")
    print("pip install python-docx docxtpl docx2pdf")
    sys.exit(1)

try:
    from PIL import Image as _PIL_Image
except Exception:
    _PIL_Image = None

# --- This block makes the script self-aware of its location ---
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR = Path(SCRIPT_DIR)

# --- This block ensures we can import config.py correctly ---
sys.path.insert(0, SCRIPT_DIR)
try:
    from config import PROVINCE_PROFILES, APP_VERSION
except ImportError:
    print("FATAL ERROR: Could not find 'config.py'. Make sure it's in the same folder as 'build.py'.")
    sys.exit(1)
finally:
    sys.path.pop(0)

def cleanup_icon_cache():
    """Clean Windows icon cache to ensure new icons are displayed."""
    try:
        # Clear icon cache
        os.system('ie4uinit.exe -ClearIconCache')
        # Rebuild icon cache
        os.system('ie4uinit.exe -show')
        print("Icon cache cleared successfully.")
    except Exception as e:
        print(f"Warning: Could not clear icon cache: {e}")

def force_remove_file(filepath):
    """Force remove a file with retries and error handling."""
    max_retries = 3
    for attempt in range(max_retries):
        try:
            if os.path.exists(filepath):
                os.remove(filepath)
                return True
        except Exception as e:
            if attempt == max_retries - 1:
                print(f"Error removing {filepath}: {e}")
                return False
            time.sleep(1)  # Wait before retry
    return False

def cleanup(temp_dir):
    """Cleans up all temporary build files and folders, plus stray spec/pyc caches.
    If deletion fails (Windows file locks/AV), move the folder under temp_build/stale/ for manual cleanup.
    """
    if not temp_dir or not os.path.exists(temp_dir):
        return

    print(f"Cleaning up temporary directory: {temp_dir}")
    try:
        # First, try to remove the directory and its contents
        shutil.rmtree(temp_dir, ignore_errors=True)
        
        # If the directory still exists, try moving it to a stale directory
        if os.path.exists(temp_dir):
            stale_dir = os.path.join(TEMP_BUILD_BASE_DIR, "stale")
            os.makedirs(stale_dir, exist_ok=True)
            
            # Generate a unique name for the stale directory
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            stale_path = os.path.join(stale_dir, f"stale_{timestamp}_{os.path.basename(temp_dir)}")
            
            try:
                shutil.move(temp_dir, stale_path)
                print(f"Moved to stale directory: {stale_path}")
            except Exception as e:
                print(f"Warning: Could not move stale directory: {e}")
    except Exception as e:
        print(f"Warning: Error during cleanup: {e}")
        
    # Also clean up any .pyc files in the project directory
    try:
        for root, dirs, files in os.walk(SCRIPT_DIR):
            for file in files:
                if file.endswith('.pyc') or file.endswith('.pyo'):
                    try:
                        os.remove(os.path.join(root, file))
                    except:
                        pass
    except Exception as e:
        print(f"Warning: Error cleaning up .pyc files: {e}")

def is_path_watched(path):
    """Check if a path is being watched by a file system filter driver (like Terabox)."""
    try:
        # First, check if the path exists and is accessible
        if not os.path.exists(path):
            return False
            
        # Try to create a test file in the directory
        test_file = os.path.join(path, f"_test_{os.getpid()}.tmp")
        try:
            with open(test_file, 'w') as f:
                f.write('test')
            os.remove(test_file)
            return False  # If we can create and delete a file, it's probably not locked
        except (IOError, OSError):
            # If we can't create a file, the directory might be watched
            return True
    except Exception as e:
        print(f"Warning: Error checking if path is watched ({path}): {e}")
        # If there's any error, assume it's not being watched to avoid false positives
        return False

# --- Configuration using project-relative paths ---
MAIN_SCRIPT_NAME = "main.py"
ICON_FILE_NAME = "logo.ico"
LOGO_FILE_NAME = "logo.png"
# --- REVISED FILENAME TEMPLATE (uses APP_VERSION) ---
APP_NAME_TEMPLATE = f"DOLE_v{APP_VERSION}_{{province}}"
# Place temp builds outside the project (avoids Terabox/backup locks)
TEMP_BUILD_BASE_DIR = os.path.join(tempfile.gettempdir(), "DOLEBuild")
DIST_DIR = str(PROJECT_DIR / "MIMAROPA")
AUDITOR_SCRIPT_NAME = "auditor.py"
AUDITOR_APP_NAME_TEMPLATE = f"DOLE_v{APP_VERSION}_{{province}}_Auditor"

# Nickname Generator (always blue theme; no province injection)
NICKNAME_GUI_SCRIPT_NAME = "nickname_generator.py"
NICKNAME_APP_NAME = f"DOLE_v{APP_VERSION}_Nickname_Generator"
NICKNAME_SOURCE_DIR = "Nickname Generator"


def force_remove_directory(dir_path):
    """Robustly removes a directory, retrying on PermissionError. Returns True on success."""
    if not os.path.exists(dir_path):
        return True
    retries, delay = 6, 1
    for i in range(retries):
        try:
            shutil.rmtree(dir_path)
            return True
        except PermissionError as e:
            if i < retries - 1:
                print(f"  Warning: Cleanup failed on attempt {i+1}. Retrying in {delay}s...")
                time.sleep(delay)
            else:
                print(f"  Warning: Could not remove directory '{dir_path}' after {retries} attempts. It may be locked by AV or OS cache.")
                print(f"  Reason: {e}")
                return False

def _province_alias(name: str) -> str:
    alias_map = {
        "Oriental Mindoro": "OrMin",
        "Occidental Mindoro": "OccMin",
        "Marinduque": "Marinduque",
        "Romblon": "Romblon",
        "Palawan": "Palawan",
    }
    return alias_map.get(name, name.replace(" ", "_"))

def _output_dir_for(province: str | None, for_nickname: bool = False) -> str:
    base = DIST_DIR
    if for_nickname or not province:
        return base
    return os.path.join(base, _province_alias(province))

def build_nickname_gui():
    """Build the blue-themed Nickname Generator GUI as a standalone EXE."""
    out_dir = _output_dir_for(None, for_nickname=True)
    os.makedirs(out_dir, exist_ok=True)
    exe_path = os.path.join(out_dir, f"{NICKNAME_APP_NAME}.exe")
    if os.path.exists(exe_path):
        try:
            os.remove(exe_path)
            print(f"Removed existing executable: {exe_path}")
        except OSError as e:
            print(f"Warning: Could not remove existing executable '{exe_path}': {e}")

    print(f"\n>>> Starting build for: {NICKNAME_APP_NAME} <<<")
    temp_directory = None
    try:
        temp_directory = create_temp_files_nickname()
        run_pyinstaller(temp_directory, NICKNAME_APP_NAME, NICKNAME_GUI_SCRIPT_NAME, province=None, for_nickname=True)
        return True
    except FileNotFoundError as e:
        print(f"\n---!!! NICKNAME GENERATOR BUILD FAILED: Missing File !!!---")
        print(f"Error: {e}")
    except Exception:
        print(f"\n---!!! AN ERROR OCCURRED DURING NICKNAME GENERATOR BUILD !!!---")
    finally:
        if temp_directory:
            cleanup(temp_directory)

    return False

def create_temp_files(province):
    """Creates a unique temporary directory with modified source files for this build run."""
    print(f"--- Preparing temporary files for {province}... ---")
    os.makedirs(TEMP_BUILD_BASE_DIR, exist_ok=True)
    unique_dir = os.path.join(TEMP_BUILD_BASE_DIR, f"session_{int(time.time())}_{os.getpid()}")
    # Ensure the path is clean
    force_remove_directory(unique_dir)
    os.makedirs(unique_dir, exist_ok=True)

    # --- MODIFIED: Added excel_converter.py to the list of source files ---
    source_files = [
        "main.py", "gui.py", "analysis_engine.py",
        "data_utils.py", "config.py", "excel_converter.py", "auditor.py"
    ]
    for file_name in source_files:
        source_path = os.path.join(SCRIPT_DIR, file_name)
        if not os.path.exists(source_path):
            raise FileNotFoundError(f"Required script file not found: {source_path}")
        shutil.copy(source_path, unique_dir)

    temp_main_path = os.path.join(unique_dir, MAIN_SCRIPT_NAME)
    with open(temp_main_path, 'r+', encoding='utf-8') as f:
        content = f.read()
        f.seek(0, 0)
        f.write(f"DEFAULT_PROVINCE = \"{province}\"\n\n" + content)
    # Inject province into auditor as well (theme/title awareness)
    temp_auditor_path = os.path.join(unique_dir, AUDITOR_SCRIPT_NAME)
    if os.path.exists(temp_auditor_path):
        with open(temp_auditor_path, 'r+', encoding='utf-8') as f:
            content = f.read()
            f.seek(0, 0)
            f.write(f"DEFAULT_PROVINCE = \"{province}\"\n\n" + content)
    
    print("Temporary files created successfully.")
    return unique_dir

def create_temp_files_nickname():
    """Create a unique temporary directory for Nickname Generator build.
    Copies only the needed sources and does NOT inject any province.
    """
    print(f"--- Preparing temporary files for Nickname Generator... ---")
    os.makedirs(TEMP_BUILD_BASE_DIR, exist_ok=True)
    unique_dir = os.path.join(TEMP_BUILD_BASE_DIR, f"session_{int(time.time())}_{os.getpid()}")
    force_remove_directory(unique_dir)
    os.makedirs(unique_dir, exist_ok=True)

    # Core sources
    base_sources = [
        NICKNAME_GUI_SCRIPT_NAME,
    ]
    for file_name in base_sources:
        source_path = os.path.join(SCRIPT_DIR, file_name)
        if not os.path.exists(source_path):
            raise FileNotFoundError(f"Required script file not found: {source_path}")
        shutil.copy(source_path, unique_dir)

    print("Temporary files (Nickname) created successfully.")
    return unique_dir

def get_build_output_path(app_name, province: str | None = None, for_nickname: bool = False):
    """Determine the best output path for the built executable under MIMAROPA tree."""
    dist_dir = _output_dir_for(province, for_nickname)
    # If the output directory is being watched, use a temp directory
    if is_path_watched(dist_dir):
        alt_dir = os.path.join(tempfile.gettempdir(), "DOLE_Builds")
        os.makedirs(alt_dir, exist_ok=True)
        print(f"\n!!! NOTICE: Standard output directory is being watched by Terabox/backup.")
        print(f"!!! Build output will be placed in: {alt_dir}")
        return os.path.join(alt_dir, f"{app_name}.exe")
    
    return os.path.join(dist_dir, f"{app_name}.exe")

def run_pyinstaller(temp_dir, app_name, script_filename, province: str | None = None, for_nickname: bool = False):
    """Executes the PyInstaller command for a given script and app name."""
    print(f"--- Starting PyInstaller build for {app_name}... ---")
    
    # Determine the best output path
    final_exe_path = get_build_output_path(app_name, province=province, for_nickname=for_nickname)
    final_dir = os.path.dirname(final_exe_path)

    # Ensure output directory exists
    os.makedirs(final_dir, exist_ok=True)

    script_path = os.path.join(temp_dir, script_filename)
    icon_path = os.path.join(SCRIPT_DIR, ICON_FILE_NAME)
    logo_path = os.path.join(SCRIPT_DIR, LOGO_FILE_NAME)
    
    # Clean up any existing executable and related files
    exe_name = f"{app_name}.exe"
    for f in [final_exe_path, 
              os.path.join(final_dir, f"{app_name}.pkg"),
              os.path.join(final_dir, f"{app_name}.spec")]:
        force_remove_file(f)
    
    # Clear icon cache before building
    cleanup_icon_cache()

    # Build into a TEMP dist folder to avoid AV/indexer locking during manifest embed
    temp_dist_dir = os.path.join(temp_dir, "dist")
    os.makedirs(temp_dist_dir, exist_ok=True)

    # Ensure output directory exists
    os.makedirs(final_dir, exist_ok=True)
    
    # Check for assets
    has_icon = os.path.exists(icon_path)
    has_logo = os.path.exists(logo_path)
    if not has_icon:
        print(f"Warning: Icon not found at {icon_path}. Proceeding without an icon.")
    if not has_logo:
        print(f"Warning: Logo not found at {logo_path}. Proceeding without embedded logo.")

    # Collect all data files to include
    data_files = []
    
    # Add logo files (do not shadow logo_path variable above)
    for _asset in ['logo.ico', 'logo.png', 'logo.jpg']:
        asset_path = os.path.join(SCRIPT_DIR, _asset)
        if os.path.exists(asset_path):
            data_files.extend(["--add-data", f"{asset_path};."])
    
    # Removed unused CSV data files from bundle to reduce size and improve startup
    
    command = [
        sys.executable, "-m", "PyInstaller", "--noconfirm",
        "--name", app_name, "--onefile", "--windowed", "--noupx",
        "--hidden-import=pandas._libs.tslibs.np_datetime",
        "--hidden-import=pandas._libs.tslibs.nattype",
        "--hidden-import=pandas._libs.tslibs.timezones",
        "--hidden-import=win32com",
        "--hidden-import=win32com.client",
        "--hidden-import=pythoncom",
        "--hidden-import=pywintypes",
        "--hidden-import=rapidfuzz",
        "--hidden-import=jellyfish",
        "--hidden-import=openpyxl",
        "--hidden-import=openpyxl.styles",
        "--hidden-import=openpyxl.worksheet",
        "--add-binary", "C:\\Windows\\System32\\vcruntime140_1.dll;.",
        "--add-binary", "C:\\Windows\\System32\\msvcp140.dll;.",
        "--add-binary", "C:\\Windows\\System32\\vcruntime140.dll;.",
        "--add-binary", "C:\\Windows\\System32\\concrt140.dll;.",
        "--runtime-hook", os.path.join(os.path.dirname(sys.executable), "Lib", "site-packages", "PyInstaller", "hooks", "rthooks", "pyi_rth_inspect.py")
    ]
    
    # Add data files if they exist
    command.extend(data_files)
    if has_icon:
        command.append(f"--icon={icon_path}")
    if has_logo:
        command.append(f"--add-data={logo_path}{os.pathsep}.")

    # Optional native splash (enabled: use bootloader splash image during unpacking)
    USE_NATIVE_SPLASH = True
    if USE_NATIVE_SPLASH:
        # Province-specific splash lookup
        splash_path = None
        candidates = []
        if province:
            # Short aliases
            alias_map = {
                "Oriental Mindoro": "OrMin",
                "Occidental Mindoro": "OccMin",
                "Marinduque": "Marinduque",
                "Romblon": "Romblon",
                "Palawan": "Palawan",
            }
            alias = alias_map.get(province)
            slug = province.replace(" ", "_")
            if alias:
                candidates.extend([f"{alias}_Splash.png", f"{alias}_Splash.jpg"])
            candidates.extend([f"{slug}_Splash.png", f"{slug}_Splash.jpg"])
        # Generic fallback (no logo fallback by request)
        candidates.extend(["splash.png", "splash.jpg"])
        for candidate in candidates:
            cand_abs = os.path.join(SCRIPT_DIR, candidate)
            if os.path.exists(cand_abs):
                splash_path = cand_abs
                break
        if splash_path:
            # Optional auto-resize to max width 600px for faster bootloader draw
            try:
                if _PIL_Image is not None:
                    with _PIL_Image.open(splash_path) as im:
                        w, h = im.size
                        if w > 600:
                            ratio = 600 / float(w)
                            new_size = (600, max(1, int(h * ratio)))
                            resized_path = os.path.join(temp_dir, f"_splash_{province or 'generic'}.png")
                            im.convert("RGBA").resize(new_size, resample=getattr(_PIL_Image, 'LANCZOS', 1)).save(resized_path)
                            splash_path = resized_path
            except Exception as _img_err:
                print(f"Warning: Could not auto-resize splash image: {_img_err}")
            command.append(f"--splash={splash_path}")
    command.extend([
        f"--distpath={temp_dist_dir}",
        f"--workpath={os.path.join(temp_dir, 'build')}",
        f"--specpath={temp_dir}",
        script_path,
    ])
    
    print(f"Executing command: {' '.join(command)}")
    try:
        # First, try without shell=True as it's more reliable
        try:
            process = subprocess.run(
                command, 
                check=True, 
                capture_output=True, 
                text=True, 
                encoding='utf-8', 
                errors='ignore'
            )
        except subprocess.CalledProcessError as e:
            # If that fails, try with shell=True as fallback
            print("Initial build failed, retrying with shell=True...")
            process = subprocess.run(
                command, 
                check=True, 
                capture_output=True, 
                text=True, 
                encoding='utf-8', 
                errors='ignore',
                shell=True
            )
            
        # Print the output
        if process.stdout:
            print("\n=== PyInstaller Output ===")
            print(process.stdout)
        if process.stderr:
            print("\n=== PyInstaller Warnings/Errors ===")
            print(process.stderr)
        # After a successful build in temp dist, move the EXE to the final location
        src_exe = os.path.join(temp_dist_dir, f"{app_name}.exe")
        dst_exe = final_exe_path

        def _move_with_retries(src, dst):
            max_retries = 10
            base_delay = 0.5
            last_err = None
            
            # Ensure destination directory exists
            os.makedirs(os.path.dirname(dst), exist_ok=True)
            
            # First try: Direct move with replace
            for attempt in range(max_retries):
                try:
                    if os.path.exists(dst):
                        force_remove_file(dst)
                    os.replace(src, dst)
                    # Clear icon cache after successful move
                    cleanup_icon_cache()
                    return True
                except (OSError, PermissionError) as e:
                    last_err = e
                    if attempt < max_retries - 1:  # Don't sleep on last attempt
                        time.sleep(base_delay * (1.5 ** attempt))  # Exponential backoff
            
            # Second try: Copy then delete
            try:
                shutil.copy2(src, dst)
                try:
                    os.remove(src)
                except OSError:
                    print(f"Warning: Could not remove source file {src} after copy")
                cleanup_icon_cache()
                return True
            except Exception as e:
                last_err = e
            
            # Third try: Move to a temp location in the same directory first
            try:
                temp_dst = f"{dst}.temp"
                if os.path.exists(temp_dst):
                    force_remove_file(temp_dst)
                shutil.copy2(src, temp_dst)
                os.rename(temp_dst, dst)
                try:
                    os.remove(src)
                except OSError:
                    print(f"Warning: Could not remove source file {src}")
                cleanup_icon_cache()
                return True
            except Exception as e:
                last_err = e
            
            # Final fallback: Try to move to a different location and inform user
            alt_dst = os.path.join(tempfile.gettempdir(), os.path.basename(dst))
            try:
                shutil.copy2(src, alt_dst)
                print(f"\n!!! WARNING: Could not move to {dst} due to file locking")
                print(f"!!! The built executable has been saved to: {alt_dst}")
                try:
                    os.remove(src)
                except OSError:
                    pass
                cleanup_icon_cache()
                return True
            except Exception as e:
                raise RuntimeError(
                    f"Failed to move built EXE after multiple attempts. "
                    f"Last error: {last_err}. "
                    f"Temporary build is at: {src}"
                )

        if not os.path.exists(src_exe):
            raise FileNotFoundError(f"Expected built EXE not found at {src_exe}")
        success = _move_with_retries(src_exe, dst_exe)
        print(f"\n--- Build successful for {app_name}! ---")
        print(f"Executable is located at: {dst_exe}")
        
        # If we're not using the standard dist directory, provide a copy command
        if not os.path.commonprefix([os.path.abspath(dst_exe), os.path.abspath(DIST_DIR)]) == os.path.abspath(DIST_DIR):
            print("\nTo copy the executable to the standard location, run:")
            print(f'copy "{dst_exe}" "{os.path.join(DIST_DIR, os.path.basename(dst_exe))}"')
            print("\nNote: You may need to close Terabox or other backup software first.")
        
        # Return True to indicate success
        return True
    except subprocess.CalledProcessError as e:
        print(f"\n--- PyInstaller failed with error code {e.returncode} ---")
        if e.stdout:
            print("\n=== STDOUT ===")
            print(e.stdout)
        if e.stderr:
            print("\n=== STDERR ===")
            print(e.stderr)
        
        # Additional debug information
        print("\n=== Additional Debug Information ===")
        print(f"Python Executable: {sys.executable}")
        print(f"Working Directory: {os.getcwd()}")
        print(f"Temporary Directory: {temp_dir}")
        print(f"Script Path: {script_path}")
        print(f"Script Exists: {os.path.exists(script_path)}")
        print(f"Icon Exists: {os.path.exists(icon_path) if has_icon else 'N/A'}")
        print(f"Logo Exists: {os.path.exists(logo_path) if has_logo else 'N/A'}")
        
        # Check for common PyInstaller issues
        print("\n=== Checking for Common Issues ===")
        return False

def copy_readme_for_province(province_name, out_dir: str):
    """Create a province-embedded README in dist as README_<Province>.txt.
    
    Uses README_template.txt from the docs folder as the source.
    """
    try:
        # Path to README template in the docs folder
        template_path = os.path.join(SCRIPT_DIR, "docs", "README_template.txt")
        
        # Create province output directory if it doesn't exist
        os.makedirs(out_dir, exist_ok=True)
        
        # Output file path
        dst_readme = os.path.join(out_dir, f"README_{province_name.replace(' ', '_')}.txt")
        
        # Check if template exists
        if not os.path.exists(template_path):
            print(f"Error: README template not found at {template_path}")
            return False
            
        # Read template and replace placeholders
        with open(template_path, 'r', encoding='utf-8') as f:
            content = f.read()
            
        # Replace placeholders
        content = content.replace('{{PROVINCE}}', province_name)
        content = content.replace('{{CURRENT_DATE}}', datetime.now().strftime('%Y-%m-%d'))
        
        # Write to destination
        with open(dst_readme, 'w', encoding='utf-8') as f:
            f.write(content)
            
        print(f"Created README for {province_name} at: {dst_readme}")
        return True
        
    except Exception as e:
        print(f"Error creating README for {province_name}: {str(e)}")
        return False

def generate_user_manual(province_name, output_dir=None):
    """
    Generate a user manual PDF for the specified province.
    Automatically overwrites existing files without confirmation.
    
    Args:
        province_name (str): Name of the province
        output_dir (str, optional): Directory to save the PDF. Defaults to None (same as executable).
        
    Returns:
        bool: True if successful, False otherwise
    """
    try:
        # Path to the template file
        template_path = os.path.join(SCRIPT_DIR, "docs", "User_Manual_Template.docx")
        
        # Check if template exists
        if not os.path.exists(template_path):
            print(f"Error: User manual template not found at {template_path}")
            return False
            
        # Set output directory
        if output_dir is None:
            output_dir = _output_dir_for(province_name)
            
        # Ensure output directory exists
        os.makedirs(output_dir, exist_ok=True)
        
        # Prepare output filenames
        base_filename = f"User_Manual_v{APP_VERSION}_{province_name.replace(' ', '_')}"
        docx_output = os.path.join(output_dir, f"{base_filename}.docx")
        pdf_output = os.path.join(output_dir, f"{base_filename}.pdf")
        
        # Remove existing files if they exist (no confirmation)
        for file_path in [docx_output, pdf_output]:
            if os.path.exists(file_path):
                try:
                    os.remove(file_path)
                except Exception as e:
                    print(f"Warning: Could not remove existing file {file_path}: {e}")
        
        # Prepare context for template
        context = {
            'VERSION': APP_VERSION,
            'PROVINCE': province_name,
            'DATE': datetime.now().strftime('%Y-%m-%d')
        }
        
        print(f"Generating user manual for {province_name}...")
        
        # First, try with docxtpl for headers/footers
        try:
            print("Processing template with docxtpl (for headers/footers)...")
            doc = DocxTemplate(template_path)
            
            # Create context with only the exact placeholders we want to support
            context = {
                'VERSION': str(APP_VERSION),
                'PROVINCE': province_name,
                'DATE': datetime.now().strftime('%Y-%m-%d')
            }
            
            # Render with docxtpl (handles headers/footers)
            doc.render(context, autoescape=True)
            
            # Save the document
            os.makedirs(os.path.dirname(docx_output), exist_ok=True)
            doc.save(docx_output)
            
            # Now process the document again to handle any placeholders in the body
            # that docxtpl might have missed (case-insensitive matching)
            from docx import Document as DocxDocument
            doc = DocxDocument(docx_output)
            
            # Replace placeholders in paragraphs
            for paragraph in doc.paragraphs:
                if '{{' in paragraph.text and '}}' in paragraph.text:
                    # Replace exact placeholders in the text
                    for placeholder, value in context.items():
                        if f'{{{{{placeholder}}}}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace(
                                f'{{{{{placeholder}}}}}', 
                                value
                            )
            
            # Replace placeholders in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if '{{' in cell.text and '}}' in cell.text:
                            # Replace exact placeholders in table cells
                            for placeholder, value in context.items():
                                if f'{{{{{placeholder}}}}}' in cell.text:
                                    cell.text = cell.text.replace(
                                        f'{{{{{placeholder}}}}}', 
                                        value
                                    )
            
            # Save the document with all replacements
            doc.save(docx_output)
            
            # Convert to PDF
            convert(docx_output, pdf_output)
            print(f"Successfully generated PDF: {pdf_output}")
            
            # Remove the intermediate DOCX file if PDF generation was successful
            try:
                os.remove(docx_output)
            except Exception as e:
                print(f"Warning: Could not remove intermediate DOCX file: {e}")
            
            return True
            
        except Exception as e:
            print(f"Error with docxtpl processing: {str(e)}")
            print("Falling back to direct processing...")
            
            # Fallback to direct processing if docxtpl fails
            try:
                doc = Document(template_path)
                
                # Process the document (same as before)
                for paragraph in doc.paragraphs:
                    if '{{' in paragraph.text and '}}' in paragraph.text:
                        placeholders = re.findall(r'\{\{\s*([^}]+)\s*\}\}', paragraph.text, re.IGNORECASE)
                        for ph in placeholders:
                            ph_upper = ph.upper().strip()
                            if ph_upper in case_insensitive_context:
                                paragraph.text = re.sub(
                                    r'\{\{\s*' + re.escape(ph) + r'\s*\}\}',
                                    case_insensitive_context[ph_upper],
                                    paragraph.text,
                                    flags=re.IGNORECASE
                                )
                
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if '{{' in cell.text and '}}' in cell.text:
                                placeholders = re.findall(r'\{\{\s*([^}]+)\s*\}\}', cell.text, re.IGNORECASE)
                                for ph in placeholders:
                                    ph_upper = ph.upper().strip()
                                    if ph_upper in case_insensitive_context:
                                        cell.text = re.sub(
                                            r'\{\{\s*' + re.escape(ph) + r'\s*\}\}',
                                            case_insensitive_context[ph_upper],
                                            cell.text,
                                            flags=re.IGNORECASE
                                        )
                
                # Save and convert
                os.makedirs(os.path.dirname(docx_output), exist_ok=True)
                doc.save(docx_output)
                convert(docx_output, pdf_output)
                
                try:
                    os.remove(docx_output)
                except Exception as e:
                    print(f"Warning: Could not remove intermediate DOCX file: {e}")
                
                return True
                
            except Exception as e2:
                print(f"Error with direct processing: {str(e2)}")
                return False
                
        except Exception as e:
            print(f"Error processing template: {str(e)}")
            if 'doc' in locals():
                print(f"Document type: {type(doc)}")
            return False
        
        print(f"Successfully generated: {pdf_output}")
        return True
        
    except Exception as e:
        print(f"Error generating user manual for {province_name}: {str(e)}")
        return False


def clear_windows_icon_cache():
    """Safely refresh Windows icon cache without stopping Explorer."""
    print("\n=== Refreshing Windows icon cache ===")
    try:
        # Method 1: Use Windows API to refresh the icon cache
        try:
            import ctypes
            ctypes.windll.shell32.SHChangeNotify(0x08000000, 0x00001000, 0, 0)
            print("Icon cache refresh requested via Windows API")
        except Exception as e:
            print(f"Warning: Could not refresh via Windows API: {e}")
            
        # Method 2: Touch the icon file to force cache update
        icon_path = os.path.join(SCRIPT_DIR, ICON_FILE_NAME)
        if os.path.exists(icon_path):
            try:
                current_time = time.time()
                os.utime(icon_path, (current_time, current_time))
                print(f"Updated timestamp for {ICON_FILE_NAME}")
            except Exception as e:
                print(f"Warning: Could not update icon timestamp: {e}")
        
        # Method 3: Use Windows built-in tool (least likely to cause issues)
        try:
            subprocess.run(['ie4uinit.exe', '-show'], 
                         stdout=subprocess.PIPE, 
                         stderr=subprocess.PIPE,
                         creationflags=subprocess.CREATE_NO_WINDOW)
            print("Icon cache refresh requested via ie4uinit")
        except Exception as e:
            print(f"Warning: Could not run ie4uinit: {e}")
            
        return True
    except Exception as e:
        print(f"Error during icon cache refresh: {e}")
        return False

def build_single_province(target_province, skip_overwrite_check=False, build_auditor=False):
    """
    Handles the build process for one province's main application.
    
    Args:
        target_province: Name of the province to build for
        skip_overwrite_check: If True, will skip the overwrite confirmation
        build_auditor: If True, will also build the auditor for this province
    
    Returns:
        bool: True on success, False on failure
    """
    out_dir = _output_dir_for(target_province)
    expected_exe = os.path.join(out_dir, f"DOLE_v{APP_VERSION}_{target_province.replace(' ', '_')}.exe")
    print(f"\n{'='*50}\nBUILDING: {target_province}\nOutput will be: {expected_exe}\n{'='*50}\n")
    
    # Clear Windows icon cache before building
    clear_windows_icon_cache()
    
    # Build the main application
    app_name = APP_NAME_TEMPLATE.format(province=target_province.replace(" ", "_"))
    
    # Verify the province exists in the configuration
    if target_province not in PROVINCE_PROFILES:
        print(f"\n!!! ERROR: Province '{target_province}' not found in PROVINCE_PROFILES !!!")
        print("Available provinces:", ", ".join(PROVINCE_PROFILES.keys()))
        return False
        
    main_app_name = APP_NAME_TEMPLATE.format(province=target_province.replace(' ', '_'))
    executable_path = get_build_output_path(main_app_name, province=target_province)
    # Also pre-clean province README to ensure template updates are reflected
    try:
        readme_path = os.path.join(out_dir, f"README_{target_province.replace(' ', '_')}.txt")
        if os.path.exists(readme_path):
            os.remove(readme_path)
            print(f"Removed existing README: {readme_path}")
    except OSError as e:
        print(f"Warning: Could not remove existing README '{readme_path}': {e}")
    print(f"\n{'='*80}")
    print(f"BUILDING: {target_province}")
    print(f"Output will be: {executable_path}")
    print(f"{'='*80}\n")

    # Check if output directory is writable
    output_dir = os.path.dirname(executable_path)
    try:
        os.makedirs(output_dir, exist_ok=True)
        test_file = os.path.join(output_dir, 'test_write.tmp')
        with open(test_file, 'w') as f:
            f.write('test')
        os.remove(test_file)
    except Exception as e:
        print(f"\n!!! ERROR: Cannot write to output directory: {output_dir}")
        print(f"Error: {e}")
        print("Please check permissions or close any programs that might be locking this directory.")
        return False

    # Auto-delete an existing exe to ensure a clean rebuild
    if os.path.exists(executable_path):
        try:
            os.remove(executable_path)
            print(f"Removed existing executable: {executable_path}")
        except OSError as e:
            print(f"Warning: Could not remove existing executable '{executable_path}': {e}")
            print("This might cause issues with the build process.")
            if not input("Continue anyway? (y/n): ").lower().startswith('y'):
                return False
    
    temp_directory = None
    try:
        print(f"\n{'='*50}")
        print("STEP 1: Creating temporary build files...")
        temp_directory = create_temp_files(target_province)
        if not temp_directory or not os.path.exists(temp_directory):
            raise RuntimeError(f"Failed to create temporary directory: {temp_directory}")
        print(f"Temporary files created at: {temp_directory}")
        
        # Verify the main script was created
        temp_script = os.path.join(temp_directory, MAIN_SCRIPT_NAME)
        if not os.path.exists(temp_script):
            raise FileNotFoundError(f"Main script not found in temp directory: {temp_script}")
        print(f"Found main script: {temp_script}")
        
        print(f"\n{'='*50}")
        print("STEP 2: Building main application...")
        if not run_pyinstaller(temp_directory, main_app_name, MAIN_SCRIPT_NAME, province=target_province):
            raise RuntimeError("PyInstaller build failed for main application")
        
        # Only build auditor if explicitly requested
        if build_auditor:
            print(f"\n{'='*50}")
            print("STEP 3: Building auditor application...")
            auditor_app_name = AUDITOR_APP_NAME_TEMPLATE.format(province=target_province.replace(' ', '_'))
            if not run_pyinstaller(temp_directory, auditor_app_name, AUDITOR_SCRIPT_NAME, province=target_province):
                raise RuntimeError("PyInstaller build failed for auditor application")
        
        # Include a province-specific README and user manual
        print(f"\n{'='*50}")
        print("STEP 4: Generating documentation...")
        copy_readme_for_province(target_province, out_dir)
        
        # Generate user manual
        print(f"\n{'='*50}")
        print("STEP 5: Generating user manual...")
        if not generate_user_manual(target_province, out_dir):
            print("Warning: Failed to generate user manual")
            # Continue with build even if manual generation fails
        
        print(f"\n{'='*50}")
        print("BUILD COMPLETED SUCCESSFULLY!")
        print(f"Main Output: {executable_path}")
        if build_auditor:
            auditor_app_name = AUDITOR_APP_NAME_TEMPLATE.format(province=target_province.replace(' ', '_'))
            print(f"Auditor Output: {os.path.join(out_dir, f'{auditor_app_name}.exe')}")
        print(f"User Manual: {os.path.join(out_dir, f'User_Manual_v{APP_VERSION}_{target_province.replace(' ', '_')}.pdf')}")
        return True
        
    except FileNotFoundError as e:
        print(f"\n!!! BUILD FAILED: Missing File !!!")
        print(f"Error: {e}")
        print("\nThis might be caused by:")
        print("1. Missing required files in the project directory")
        print("2. Insufficient permissions to access files")
        print("3. Files being locked by another process")
    except Exception as e:
        print(f"\n!!! BUILD FAILED FOR {target_province} !!!")
        print(f"Error type: {type(e).__name__}")
        print(f"Error details: {str(e)}")
        print("\nStack trace:")
        traceback.print_exc()
    finally:
        if temp_directory and os.path.exists(temp_directory):
            print(f"\nCleaning up temporary files...")
            cleanup(temp_directory)
            
    return False

def build_auditor_single_province(target_province):
    """
    Build only the Auditor EXE for one province.
    Note: This function does NOT create or delete README files;
    README generation is handled by main app builds only.
    Returns True on success, False on failure.
    """
    out_dir = _output_dir_for(target_province)
    auditor_app_name = AUDITOR_APP_NAME_TEMPLATE.format(province=target_province.replace(' ', '_'))
    auditor_executable_path = os.path.join(out_dir, f"{auditor_app_name}.exe")

    # Auto-delete an existing exe to ensure a clean rebuild
    if os.path.exists(auditor_executable_path):
        try:
            os.remove(auditor_executable_path)
            print(f"Removed existing executable: {auditor_executable_path}")
        except OSError as e:
            print(f"Warning: Could not remove existing executable '{auditor_executable_path}': {e}")
            # Continue anyway; PyInstaller will overwrite if possible

    print(f"\n>>> Starting AUDITOR-ONLY build for: {target_province} <<<")

    temp_directory = None
    try:
        temp_directory = create_temp_files(target_province)
        run_pyinstaller(temp_directory, auditor_app_name, AUDITOR_SCRIPT_NAME, province=target_province)
        return True
    except FileNotFoundError as e:
        print(f"\n---!!! AUDITOR BUILD FAILED: Missing File !!!---")
        print(f"Error: {e}")
    except Exception:
        print(f"\n---!!! AN ERROR OCCURRED DURING AUDITOR BUILD FOR {target_province} !!!---")
    finally:
        if temp_directory:
            cleanup(temp_directory)

    return False

def main():
    """Main function to run the build process in a loop or non-interactively."""
    available_provinces = list(PROVINCE_PROFILES.keys())

    # Check for command-line arguments for non-interactive mode
    if len(sys.argv) > 1:
        try:
            choice = int(sys.argv[1])
            if 1 <= choice <= len(available_provinces):
                target_province = available_provinces[choice - 1]
                build_single_province(target_province, skip_overwrite_check=True)
                return
            elif choice == 6:  # Build ALL
                successful_builds = 0
                failed_builds = []
                for province in available_provinces:
                    if build_single_province(province, skip_overwrite_check=True, build_auditor=True):
                        successful_builds += 1
                    else:
                        failed_builds.append(province)
                print("\n========== 'BUILD ALL' COMPLETE ==========")
                print(f"Successfully built: {successful_builds} executables.")
                if failed_builds:
                    print(f"Failed builds: {len(failed_builds)} ({', '.join(failed_builds)})")
                print("========================================\n")
                return
            elif choice == nickname_option:  # Build Nickname Generator
                build_nickname_gui()
                return
        except ValueError:
            pass

    # Interactive mode
    while True:
        print("--------------------------------------------------")
        print("Available Provinces for Build:")
        for i, p in enumerate(available_provinces, 1):
            print(f"  {i}. {p}")
        build_all_option = len(available_provinces) + 1
        auditor_one_option = len(available_provinces) + 2
        auditor_all_option = len(available_provinces) + 3
        user_manual_one_option = len(available_provinces) + 4
        user_manual_all_option = len(available_provinces) + 5
        nickname_option = len(available_provinces) + 6
        exit_option = len(available_provinces) + 7
        print(f"  {build_all_option}. Build ALL (Main + Auditor + User Manual)")
        print(f"  {auditor_one_option}. Auditor ONLY - Choose Province")
        print(f"  {auditor_all_option}. Auditor ONLY - Build ALL Provinces")
        print(f"  {user_manual_one_option}. User Manual - Choose Province")
        print(f"  {user_manual_all_option}. User Manual - Build ALL Provinces")
        print(f"  {nickname_option}. Build Nickname Generator (Blue Theme)")
        print(f"  {exit_option}. Exit")
        print("--------------------------------------------------")

        choice_str = input(f"Enter your choice (1-{exit_option}, or Q to quit): ").strip()

        try:
            # Check for 'Q' or 'q' to quit
            if choice_str.lower() == 'q':
                print("Exiting build script. Goodbye!")
                break
                
            choice = int(choice_str)

            if choice == exit_option:
                print("Exiting build script. Goodbye!")
                break
            
            elif choice == build_all_option:
                print("\nWARNING: This will build an executable for all provinces.")
                print("Existing files will be overwritten automatically.")
                
                successful_builds = 0
                failed_builds = []
                
                for province in available_provinces:
                    print(f"\n{'='*80}")
                    print(f"PROCESSING: {province}")
                    print(f"{'='*80}")
                    
                    # Build the main application and auditor
                    success = True
                    if not build_single_province(province, skip_overwrite_check=True, build_auditor=True):
                        success = False
                        print(f"Warning: Build failed for {province}")
                    
                    # Generate user manual (always attempt, even if build failed)
                    if not generate_user_manual(province):
                        print(f"Warning: Failed to generate user manual for {province}")
                    
                    if success:
                        successful_builds += 1
                    else:
                        failed_builds.append(province)
                
                print("\n========== 'BUILD ALL' COMPLETE ==========")
                print(f"Successfully built: {successful_builds} executables.")
                if failed_builds:
                    print(f"Failed builds: {len(failed_builds)} ({', '.join(failed_builds)})")
                print("========================================\n")

            elif choice == auditor_one_option:
                print("\n--- AUDITOR ONLY: Choose a province ---")
                for i, p in enumerate(available_provinces, 1):
                    print(f"  {i}. {p}")
                sel_str = input(f"Enter province number (1-{len(available_provinces)}) or 0 to cancel: ")
                try:
                    sel = int(sel_str)
                    if sel == 0:
                        print("Cancelled.\n")
                    elif 1 <= sel <= len(available_provinces):
                        target_province = available_provinces[sel - 1]
                        build_auditor_single_province(target_province)
                    else:
                        print("Invalid selection.\n")
                except ValueError:
                    print("Invalid input.\n")

            elif choice == auditor_all_option:
                print("\nWARNING: This will build AUDITOR-ONLY executables for all provinces.")
                confirm = input("This may overwrite existing files. Continue? (y/n): ").lower().strip()
                if confirm == 'y':
                    successful = 0
                    failed = []
                    for province in available_provinces:
                        if build_auditor_single_province(province):
                            successful += 1
                        else:
                            failed.append(province)
                    print("\n====== 'AUDITOR-ONLY BUILD ALL' COMPLETE ======")
                    print(f"Successfully built auditor: {successful} executables.")
                    if failed:
                        print(f"Failed: {len(failed)} ({', '.join(failed)})")
                    print("==============================================\n")
                else:
                    print("Auditor-only build all operation cancelled.\n")

            elif choice == nickname_option:
                build_nickname_gui()
                
            # User Manual - Choose Province
            elif choice == user_manual_one_option:
                print("\n--- USER MANUAL: Choose a province ---")
                for i, p in enumerate(available_provinces, 1):
                    print(f"  {i}. {p}")
                sel_str = input(f"Enter province number (1-{len(available_provinces)}) or 0 to cancel: ")
                try:
                    sel = int(sel_str)
                    if sel == 0:
                        print("Cancelled.\n")
                    elif 1 <= sel <= len(available_provinces):
                        target_province = available_provinces[sel - 1]
                        generate_user_manual(target_province)
                    else:
                        print("Invalid selection.\n")
                except ValueError:
                    print("Invalid input.\n")

            # User Manual - Build All Provinces
            elif choice == user_manual_all_option:
                print("\nWARNING: This will generate user manuals for ALL provinces.")
                confirm = input("This may overwrite existing files. Continue? (y/n): ").lower().strip()
                if confirm == 'y':
                    successful = 0
                    failed = []
                    for province in available_provinces:
                        if generate_user_manual(province):
                            successful += 1
                        else:
                            failed.append(province)
                    print("\n====== 'USER MANUAL BUILD ALL' COMPLETE ======")
                    print(f"Successfully generated: {successful} user manuals.")
                    if failed:
                        print(f"Failed: {len(failed)} ({', '.join(failed)})")
                    print("==============================================\n")
                else:
                    print("User manual generation cancelled.\n")

            elif 1 <= choice <= len(available_provinces):
                target_province = available_provinces[choice - 1]
                build_single_province(target_province)

            else:
                print(f"Invalid number '{choice_str}'. Please enter a number between 1 and {exit_option}.\n")

        except ValueError:
            print(f"Invalid input '{choice_str}'. Please enter a number or 'Q' to quit.\n")


if __name__ == "__main__":
    main()